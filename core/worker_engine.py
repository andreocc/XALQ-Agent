import os
import datetime
import logging
import json
import requests
import re
import platform
import google.generativeai as genai
from docx import Document
from PySide6.QtCore import QSettings
from functools import lru_cache
from core.updater import Updater

# Load .env file if present
try:
    from dotenv import load_dotenv
    load_dotenv(os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '.env'))
except ImportError:
    pass  # dotenv not installed, rely on system env vars

class WorkerEngine:
    def __init__(self, base_dir=None, progress_callback=None, api_key=None):
        self.base_dir = base_dir or os.path.dirname(os.path.abspath(__file__))
        
        # Adjust base_dir if it's inside 'core'
        if os.path.basename(self.base_dir) == 'core':
            self.base_dir = os.path.dirname(self.base_dir)
            
        self.processing_dir = os.path.join(self.base_dir, 'processing')
        self.output_dir = os.path.join(self.base_dir, 'output')
        self.prompts_dir = os.path.join(self.base_dir, 'prompts')
        self.templates_dir = os.path.join(self.base_dir, 'templates')
        self.error_dir = os.path.join(self.base_dir, 'error')
        self.log_dir = os.path.join(self.base_dir, 'logs')
        
        self.progress_callback = progress_callback
        
        self._ensure_dirs()
        self.logger = self._setup_logging()
        
        self.settings = QSettings("XALQ", "XALQ Agent")
        self.updater = Updater(self.base_dir)
        
        # GitHub Config
        self.repo_url = "https://raw.githubusercontent.com/andreocc/XALQ-Agent/main/prompts/"
        self.github_pat = os.environ.get("GITHUB_PAT") or self.settings.value("github_pat", "")

        # API Key: .env -> QSettings -> empty
        self.api_key = api_key
        if not self.api_key:
            self.api_key = os.environ.get("GEMINI_API_KEY", "")
        if not self.api_key:
            self.api_key = self.settings.value("gemini_api_key", "")
        if not self.api_key:
            self.log_and_progress("Nenhuma chave de API configurada. Configure em .env ou Configurações.", "error")

        self._configure_gemini()

    def _ensure_dirs(self):
        for d in [self.processing_dir, self.output_dir, self.prompts_dir, 
                  self.templates_dir, self.error_dir, self.log_dir]:
            os.makedirs(d, exist_ok=True)

    def _setup_logging(self):
        logger = logging.getLogger("WorkerEngine")
        logger.setLevel(logging.INFO)
        # Clear existing handlers to avoid duplication if re-instantiated
        if logger.handlers:
            logger.handlers.clear()
            
        fh = logging.FileHandler(os.path.join(self.log_dir, 'worker.log'), encoding='utf-8')
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        fh.setFormatter(formatter)
        logger.addHandler(fh)
        return logger

    def _configure_gemini(self):
        if self.api_key:
            genai.configure(api_key=self.api_key)
            self.log_and_progress("Gemini configurado com sucesso.", "debug")
        else:
            self.log_and_progress("Gemini API Key não configurada!", "error")

    def log_and_progress(self, message, status_type="info"):
        if status_type == "info":
            self.logger.info(message)
        elif status_type == "error":
            self.logger.error(message)
        elif status_type == "debug":
            self.logger.debug(message)
            
        if self.progress_callback:
            self.progress_callback(message)

    @lru_cache(maxsize=32)
    def fetch_github_prompt(self, filename):
        """Fetches prompt from GitHub with caching and PAT authentication."""
        try:
            url = f"{self.repo_url}{filename}"
            headers = {}
            if self.github_pat:
                headers["Authorization"] = f"token {self.github_pat}"
            
            response = requests.get(url, headers=headers, timeout=10)
            if response.status_code == 200:
                self.log_and_progress(f"Prompt baixado do GitHub: {filename}", "debug")
                return response.text
            else:
                self.log_and_progress(f"Falha ao baixar prompt do GitHub ({response.status_code}): {filename}", "error")
                return None
        except Exception as e:
            self.log_and_progress(f"Erro de conexão GitHub: {e}", "error")
            return None

    def save_prompt_content(self, filename, content):
        try:
            path = os.path.join(self.prompts_dir, filename)
            with open(path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            self.logger.error(f"Erro ao salvar prompt {filename}: {e}")
            return False

    def read_prompt_content(self, filename):
        try:
            full_path = os.path.join(self.prompts_dir, filename)
            with open(full_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            self.log_and_progress(f"Error reading prompt file {filename}: {e}", "error")
            return None

    def load_agent_prompt(self, agent_type):
        """
        Loads prompt content. Prioritizes local file, falls back to GitHub.
        Normalizes filenames for robustness.
        """
        import unicodedata

        def normalize(text):
            text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('utf-8')
            return text.lower().replace(" ", "").replace("-", "").replace("_", "")

        target_normalized = normalize(agent_type)
        
        # 1. Try Local
        try:
            local_files = [f for f in os.listdir(self.prompts_dir) if f.endswith('.md')]
            for fname in local_files:
                # Exact
                if fname == agent_type or fname == f"{agent_type}.md":
                    return self.read_prompt_content(fname)
                # Fuzzy
                fname_normalized = normalize(fname.replace(".md", ""))
                if fname_normalized == target_normalized:
                    return self.read_prompt_content(fname)
        except Exception as e:
            self.log_and_progress(f"Error listing local prompts: {e}", "error")

        # 2. Mapped fallback
        agent_type_mapping = {
            'b2b (vende para outras empresas)': 'revenue',
            'revenue': '1_diagnostico_revenue_decision_core.md',
            'operations': '1_diagnostico_digital_operations_core.md',
        }
        
        mapped_name = agent_type_mapping.get(agent_type, agent_type)
        if not mapped_name.endswith('.md'):
             mapped_name += '.md'

        # Check local mapped
        full_path = os.path.join(self.prompts_dir, mapped_name)
        if os.path.exists(full_path):
             return self.read_prompt_content(mapped_name)

        # 3. GitHub Fetch
        self.log_and_progress(f"Prompt {mapped_name} não encontrado localmente. Buscando no GitHub...", "debug")
        content = self.fetch_github_prompt(mapped_name)
        if content:
            self.save_prompt_content(mapped_name, content)
            return content
            
        self.log_and_progress(f"Prompt not found: {mapped_name}", "error")
        return None

    def call_ai_api(self, prompt_content, config):
        # Strategy:
        # 1. Primary: User-selected model
        # 2. Fallback chain of current available models
        # NOTE: gemini-1.5-pro and gemini-1.5-flash are DEPRECATED (404).
        
        user_model = config.get('model', 'gemini-3-pro-preview')
        
        candidate_models = [
            user_model,
            'models/' + user_model if not user_model.startswith('models/') else user_model,
            
            # Current models: Pro first, Flash fallback
            'gemini-3-pro-preview',
            'gemini-2.5-pro',
            'gemini-2.5-flash',
            'gemini-2.0-flash',
            'gemini-flash-latest',
        ]
        
        # Deduplicate preserving order
        models_to_try = list(dict.fromkeys(candidate_models))
        
        last_error = None
        
        for model_name in models_to_try:
            try:
                self.log_and_progress(f"Tentando modelo: {model_name}...", "debug")
                model = genai.GenerativeModel(model_name)
                
                # Temperature Strategy: Pro = 0.1 (Precision), Flash = 0.2 (Creative/Fast)
                is_pro = "pro" in model_name.lower()
                temp = 0.1 if is_pro else 0.2
                
                generation_config = genai.types.GenerationConfig(
                    temperature=config.get('temperature', temp),
                    top_p=0.9,
                    max_output_tokens=8192,
                )
                
                response = model.generate_content(prompt_content, generation_config=generation_config)
                
                if not response.parts:
                    if response.prompt_feedback:
                         self.log_and_progress(f"Safety Block ({model_name}): {response.prompt_feedback}", "error")
                    return None # Stop on safety block usually, or continue? usually stop.

                return response.text
                
            except Exception as e:
                # Log usage limits or 404s
                self.log_and_progress(f"Erro em {model_name}: {e}", "debug")
                last_error = e
                continue
        
        self.log_and_progress(f"FALHA FATAL: Nenhum modelo disponível. Erro: {last_error}", "error")
        return None

    def parse_response(self, ai_response):
        parsed_data = {}
        # All 14 sections matching template_xalq.docx
        sections = [
            "RESUMO_EXECUTIVO", "DIAGNOSTICO", "LACUNAS", "CLASSIFICACAO",
            "ESTRUTURA_TO_BE", "MATRIZ_DE_METRICAS", "ARQUITETURA_CONCEITUAL_DE_DADOS",
            "PERGUNTAS_DECISORIAS", "KPIS_ASSOCIADOS", "VISUALIZACAO_CONCEITUAL",
            "RISCOS_ATUAIS", "RISCOS_SE_NAO_IMPLEMENTAR",
            "OBSERVACOES_XALQ", "PROXIMOS_PASSOS"
        ]
        
        # 1. Strict: [SECTION]...[/SECTION]
        for section in sections:
            pattern = fr"\[{section}\](.*?)\[/{section}\]"
            match = re.search(pattern, ai_response, re.DOTALL | re.IGNORECASE)
            parsed_data[section] = match.group(1).strip() if match else ""
            
        # 2. Flexible Fallback: header-based parsing
        if sum(1 for v in parsed_data.values() if v) < 3:
            self.log_and_progress("Parsing estrito falhou. Usando modo flexível...", "debug")
            clean_response = ai_response.replace("*", "").replace("#", "")
            current_section = None
            
            map_headers = {
                "RESUMO EXECUTIVO": "RESUMO_EXECUTIVO",
                "DIAGNOSTICO": "DIAGNOSTICO",
                "LACUNAS": "LACUNAS",
                "CLASSIFICACAO": "CLASSIFICACAO",
                "ESTRUTURA TO BE": "ESTRUTURA_TO_BE",
                "ESTRUTURA TO-BE": "ESTRUTURA_TO_BE",
                "MATRIZ DE METRICAS": "MATRIZ_DE_METRICAS",
                "ARQUITETURA CONCEITUAL": "ARQUITETURA_CONCEITUAL_DE_DADOS",
                "PERGUNTAS DECISORIAS": "PERGUNTAS_DECISORIAS",
                "KPIS ASSOCIADOS": "KPIS_ASSOCIADOS",
                "VISUALIZACAO CONCEITUAL": "VISUALIZACAO_CONCEITUAL",
                "RISCOS ATUAIS": "RISCOS_ATUAIS",
                "RISCOS SE NAO": "RISCOS_SE_NAO_IMPLEMENTAR",
                "OBSERVACOES": "OBSERVACOES_XALQ",
                "PROXIMOS PASSOS": "PROXIMOS_PASSOS",
            }
            
            lines = clean_response.split('\n')
            for line in lines:
                upper_line = line.strip().upper().replace(":", "").replace("\u00c1", "A").replace("\u00c9", "E").replace("\u00d3", "O").replace("\u00ca", "E")
                found = False
                for k, v in map_headers.items():
                    if k in upper_line and len(upper_line) < 50:
                        current_section = v
                        found = True
                        break
                
                if found: continue
                
                if current_section:
                    parsed_data[current_section] = parsed_data.get(current_section, "") + line + "\n"

        return parsed_data

    # Helper for filename
    def sanitize_filename(self, text):
        if not isinstance(text, str): return "unknown"
        s = re.sub(r'[^\w\-\.]', '_', text)
        return s.strip('_')

    def generate_word_report(self, parsed_data, agent_type, model_name, timestamp, prefix):
        template_path = os.path.join(self.templates_dir, 'template_xalq.docx')
        if not os.path.exists(template_path):
            self.log_and_progress(f"Template not found: {template_path}", "error")
            return None
            
        try:
            doc = Document(template_path)
            replacements = {
                '{{RESUMO_EXECUTIVO}}': parsed_data.get('RESUMO_EXECUTIVO', ''),
                '{{DIAGNOSTICO}}': parsed_data.get('DIAGNOSTICO', ''),
                '{{LACUNAS}}': parsed_data.get('LACUNAS', ''),
                '{{CLASSIFICACAO}}': parsed_data.get('CLASSIFICACAO', ''),
                '{{ESTRUTURA_TO_BE}}': parsed_data.get('ESTRUTURA_TO_BE', ''),
                '{{MATRIZ_DE_METRICAS}}': parsed_data.get('MATRIZ_DE_METRICAS', ''),
                '{{ARQUITETURA_CONCEITUAL_DE_DADOS}}': parsed_data.get('ARQUITETURA_CONCEITUAL_DE_DADOS', ''),
                '{{PERGUNTAS_DECISORIAS}}': parsed_data.get('PERGUNTAS_DECISORIAS', ''),
                '{{KPIS_ASSOCIADOS}}': parsed_data.get('KPIS_ASSOCIADOS', ''),
                '{{VISUALIZACAO_CONCEITUAL}}': parsed_data.get('VISUALIZACAO_CONCEITUAL', ''),
                '{{RISCOS_ATUAIS}}': parsed_data.get('RISCOS_ATUAIS', ''),
                '{{RISCOS_SE_NAO_IMPLEMENTAR}}': parsed_data.get('RISCOS_SE_NAO_IMPLEMENTAR', ''),
                '{{OBSERVACOES_XALQ}}': parsed_data.get('OBSERVACOES_XALQ', ''),
                '{{PROXIMOS_PASSOS}}': parsed_data.get('PROXIMOS_PASSOS', ''),
                '{{tipo_agente}}': agent_type,
                '{{modelo_gemini}}': model_name,
                '{{timestamp}}': timestamp
            }
            
            # Helper to replace in paragraphs
            def replace_in_p(p):
                for key, val in replacements.items():
                    if key in p.text:
                        p.text = p.text.replace(key, str(val))
            
            for p in doc.paragraphs: replace_in_p(p)
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs: replace_in_p(p)

            out_name = f"{self.sanitize_filename(prefix)}_{self.sanitize_filename(model_name)}_report_{timestamp}.docx"
            out_path = os.path.join(self.output_dir, out_name)
            doc.save(out_path)
            return out_path
        except Exception as e:
            self.log_and_progress(f"Erro na geração do DOCX: {e}", "error")
            return None

    def get_available_models(self):
        # Return sensible defaults, dynamic fetched in UI if needed, 
        # but here we just need to return what's "supported" by the tool
        return ["gemini-1.5-pro", "gemini-2.0-flash", "gemini-1.5-flash"]

    def get_prompts_list(self):
        try:
             # Just list local ones + hardcoded knowns
             local = [os.path.splitext(f)[0] for f in os.listdir(self.prompts_dir) if f.endswith('.md')]
             return list(set(local + ["revenue", "operations"]))
        except:
             return ["revenue", "operations"]
             
    def load_data(self, file_path):
        import pandas as pd
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Simple validation
            if df.empty:
                return None, []
                
            # Create a list of "Index: CompanyName" for combo box
            # Improved heuristic for Company Name
            company_col = None
            candidates = ['nome da empresa', 'empresa', 'company', 'name', 'cliente', 'organization', 'razão social', 'razao social']
            
            # 1. Search for exact/partial match
            for col in df.columns:
                if any(c in str(col).lower() for c in candidates):
                    company_col = col
                    break
            
            # 2. Fallback: First column that is NOT a timestamp
            if not company_col:
                for col in df.columns:
                    col_str = str(col).lower()
                    # Check for date keywords
                    if any(x in col_str for x in ['data', 'date', 'time', 'carimbo', 'timestamp', 'hora']):
                        continue
                    company_col = col
                    break
            
            # 3. Final Fallback
            if not company_col:
                company_col = df.columns[0]
            
            items = []
            for idx, row in df.iterrows():
                val = str(row[company_col])
                items.append(f"{idx}: {val}")
                
            return df, items
        except Exception as e:
            self.log_and_progress(f"Erro ao carregar dados: {e}", "error")
            return None, []

    def process_file(self, file_path, model_override=None, rows_to_process=None, prompt_type_override=None):
        import pandas as pd
        self.log_and_progress(f"Lendo arquivo: {file_path}")
        
        df, items = self.load_data(file_path)
        if df is None: return []

        generated_files = []
        
        # User Defined Defaults
        config = {
            'model': model_override or 'gemini-1.5-pro',
            'temperature': 0.1 # Base temp, call_ai_api adjusts per model
        }

        total = len(df)
        self.log_and_progress(f"Iniciando processamento de {total} linhas...")

        for row_idx, row in df.iterrows():
            if rows_to_process and row_idx not in rows_to_process: continue
            
            # Name Prefix
            prefix = f"Row_{row_idx}"
            
            # Re-detect col for prefix (redundant but safe)
            for col in df.columns:
                 if any(c in str(col).lower() for c in ['nome da empresa', 'empresa', 'company']):
                      prefix = str(row[col])
                      break
            
            self.log_and_progress(f"--- Processando: {prefix} ({row_idx+1}/{total}) ---")
            
            # 1. Determine Prompt
            if prompt_type_override and "Automático" not in prompt_type_override:
                p_type = prompt_type_override
            else:
                 # Auto detect
                 # Heuristic: Find a column asking for "modelo de atuação"
                 p_type = "revenue" # Default backup
                 for c in df.columns:
                     if "modelo" in str(c).lower():
                         p_type = str(row[c]).strip()
                         break
            
            # 2. Load Prompt
            prompt_text = self.load_agent_prompt(p_type)
            if not prompt_text:
                self.log_and_progress(f"Prompt não encontrado para '{p_type}'. Pulando.", "error")
                continue
                
            # 3. Call AI
            full_prompt = f"{prompt_text}\n\nDADOS DO CLIENTE:\n{row.to_string()}"
            response = self.call_ai_api(full_prompt, config)
            
            if not response:
                self.log_and_progress("Falha na geração da IA.", "error")
                continue
                
            # 4. Parse & Save
            parsed = self.parse_response(response)
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            rpt = self.generate_word_report(parsed, p_type, config['model'], timestamp, prefix)
            
            if rpt:
                generated_files.append(rpt)
                self.log_and_progress("Relatório gerado com sucesso.", "info")
        
        self.log_and_progress("Processamento finalizado.")
        return generated_files
