import re
from docx import Document

d = Document(r'c:\Dev\Xalq_Agent\templates\template_xalq.docx')
text = '\n'.join([p.text for p in d.paragraphs])
placeholders = sorted(set(re.findall(r'\{\{(\w+)\}\}', text)))

with open(r'c:\Dev\Xalq_Agent\template_analysis.txt', 'w', encoding='utf-8') as f:
    f.write('=== PLACEHOLDERS ===\n')
    for p in placeholders:
        f.write(f'  {p}\n')
    f.write(f'\nTotal: {len(placeholders)}\n')
    f.write('\n=== FULL TEXT ===\n')
    for i, p in enumerate(d.paragraphs):
        if p.text.strip():
            f.write(f'{i}: {p.text}\n')

print('Done. See template_analysis.txt')
