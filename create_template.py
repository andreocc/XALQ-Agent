from docx import Document
import os

def create_template_docx(file_path):
    document = Document()
    document.add_heading('Relatório XALQ', level=1)

    document.add_heading('Resumo Executivo', level=2)
    document.add_paragraph('{{RESUMO_EXECUTIVO}}')

    document.add_heading('Diagnóstico', level=2)
    document.add_paragraph('{{DIAGNOSTICO}}')

    document.add_heading('Lacunas', level=2)
    document.add_heading('Lacunas', level=2)
    document.add_paragraph('{{LACUNAS}}')

    document.add_heading('Classificação', level=2)
    document.add_paragraph('{{CLASSIFICACAO}}')

    document.add_heading('Observações XALQ', level=2)
    document.add_paragraph('{{OBSERVACOES_XALQ}}')

    document.add_paragraph(f'Tipo de Agente: {{tipo_agente}}')
    document.add_paragraph(f'Modelo Ollama: {{modelo_ollama}}')
    document.add_paragraph(f'Data/Hora da Geração: {{timestamp}}')

    document.save(file_path)
    print(f"Created basic template: {file_path}")

if __name__ == "__main__":
    template_dir = 'templates'
    template_file = os.path.join(template_dir, 'template_xalq.docx')
    os.makedirs(template_dir, exist_ok=True) # Ensure directory exists
    create_template_docx(template_file)
